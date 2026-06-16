#!/usr/bin/env python3
from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


DEFAULT_TITLE = "Manuscript"
DEFAULT_AUTHOR = ""
DEFAULT_OUTPUT = "vellum-ready.docx"

HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")
ORDERED_RE = re.compile(r"^\s*\d+\.\s+")
UNORDERED_RE = re.compile(r"^\s*[-*+]\s+")
FENCE_RE = re.compile(r"^\s*```")
HR_RE = re.compile(r"^\s*(?:---|\*\*\*|___)\s*$")
INLINE_RE = re.compile(r"(\*\*\*.+?\*\*\*|\*\*.+?\*\*|\*.+?\*|`.+?`)")


def parse_args():
    parser = argparse.ArgumentParser(
        description="Convert Markdown manuscripts into Vellum-friendly DOCX files."
    )
    parser.add_argument("inputs", nargs="+", help="Markdown input file(s)")
    parser.add_argument(
        "-o",
        "--output",
        default=DEFAULT_OUTPUT,
        help=f"Output DOCX path (default: {DEFAULT_OUTPUT})",
    )
    parser.add_argument(
        "--title",
        default=DEFAULT_TITLE,
        help="Document title stored in metadata (default: Manuscript)",
    )
    parser.add_argument(
        "--author",
        default=DEFAULT_AUTHOR,
        help="Document author stored in metadata",
    )
    parser.add_argument(
        "--chapter-level",
        type=int,
        default=2,
        choices=range(1, 7),
        metavar="N",
        help="Markdown heading level that starts a new chapter (default: 2 for ##)",
    )
    return parser.parse_args()


def configure_document(doc: Document, title: str, author: str):
    props = doc.core_properties
    props.title = title
    props.author = author
    props.subject = "Vellum-friendly Markdown manuscript"
    props.keywords = "Vellum, DOCX, Markdown, manuscript"

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Georgia"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name in ("Heading 1", "Heading 2", "Heading 3", "Heading 4", "Heading 5", "Heading 6"):
        style = styles[name]
        style.font.name = "Georgia"

    h1 = styles["Heading 1"]
    h1.font.size = Pt(14)
    h1.font.bold = True
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(8)

    h2 = styles["Heading 2"]
    h2.font.size = Pt(12)
    h2.font.bold = True
    h2.paragraph_format.space_before = Pt(10)
    h2.paragraph_format.space_after = Pt(4)


def add_inline_runs(paragraph, text: str):
    pos = 0
    for match in INLINE_RE.finditer(text):
        if match.start() > pos:
            paragraph.add_run(text[pos:match.start()])
        token = match.group(0)
        if token.startswith("```"):
            run = paragraph.add_run(token[3:-3])
            run.font.name = "Courier New"
        elif token.startswith("***"):
            run = paragraph.add_run(token[3:-3])
            run.bold = True
            run.italic = True
        elif token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("*"):
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Courier New"
        pos = match.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def flush_paragraph(doc: Document, buffer: list[str]):
    text = " ".join(part.strip() for part in buffer).strip()
    buffer.clear()
    if not text:
        return
    p = doc.add_paragraph(style="Normal")
    add_inline_runs(p, text)


def add_quote(doc: Document, lines: list[str]):
    text = " ".join(line.strip() for line in lines).strip()
    if not text:
        return
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.left_indent = Inches(0.6)
    p.paragraph_format.right_indent = Inches(0.6)
    p.paragraph_format.space_after = Pt(4)
    add_inline_runs(p, text)
    if p.runs:
        for run in p.runs:
            run.italic = True


def add_code_block(doc: Document, lines: list[str]):
    if not lines:
        return
    for line in lines:
        p = doc.add_paragraph(style="Normal")
        p.paragraph_format.left_indent = Inches(0.4)
        p.paragraph_format.right_indent = Inches(0.4)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line.rstrip())
        run.font.name = "Courier New"


def add_list_item(doc: Document, text: str, ordered: bool, level: int = 0):
    style = "List Number" if ordered else "List Bullet"
    p = doc.add_paragraph(style=style)
    if level:
        p.paragraph_format.left_indent = Inches(0.3 * level)
    add_inline_runs(p, text)


def add_heading(doc: Document, level: int, text: str, chapter_level: int, seen_chapter: bool):
    if level >= chapter_level:
        if seen_chapter:
            doc.add_page_break()
        p = doc.add_paragraph(style=f"Heading {min(level - chapter_level + 1, 6)}")
        p.add_run(text)
        return True

    if level == 1:
        p = doc.add_paragraph(style="Heading 1")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(text)
    else:
        style_name = f"Heading {min(level + 1, 6)}"
        p = doc.add_paragraph(style=style_name)
        p.add_run(text)
    return seen_chapter


def process_markdown(doc: Document, path: Path, chapter_level: int, first_file: bool, seen_chapter: bool):
    if not first_file:
        doc.add_page_break()

    lines = path.read_text(encoding="utf-8").splitlines()
    paragraph_buffer: list[str] = []
    quote_buffer: list[str] = []
    code_buffer: list[str] = []
    list_buffer: list[str] = []
    list_ordered = False
    in_code = False

    def flush_all():
        nonlocal quote_buffer, code_buffer, list_buffer, list_ordered
        flush_paragraph(doc, paragraph_buffer)
        if quote_buffer:
            add_quote(doc, quote_buffer)
            quote_buffer = []
        if code_buffer:
            add_code_block(doc, code_buffer)
            code_buffer = []
        if list_buffer:
            for item in list_buffer:
                add_list_item(doc, item, list_ordered)
            list_buffer = []

    for raw in lines:
        line = raw.rstrip()

        if FENCE_RE.match(line):
            flush_paragraph(doc, paragraph_buffer)
            if in_code:
                add_code_block(doc, code_buffer)
                code_buffer = []
                in_code = False
            else:
                if quote_buffer:
                    add_quote(doc, quote_buffer)
                    quote_buffer = []
                if list_buffer:
                    for item in list_buffer:
                        add_list_item(doc, item, list_ordered)
                    list_buffer = []
                in_code = True
            continue

        if in_code:
            code_buffer.append(line)
            continue

        if not line.strip():
            flush_all()
            continue

        heading = HEADING_RE.match(line)
        if heading:
            flush_all()
            hashes, text = heading.groups()
            seen_chapter = add_heading(doc, len(hashes), text.strip(), chapter_level, seen_chapter)
            continue

        if HR_RE.match(line):
            flush_all()
            doc.add_paragraph(" ", style="Normal")
            continue

        if line.lstrip().startswith(">"):
            flush_paragraph(doc, paragraph_buffer)
            if list_buffer:
                for item in list_buffer:
                    add_list_item(doc, item, list_ordered)
                list_buffer = []
            quote_buffer.append(line.lstrip()[1:].lstrip())
            continue

        unordered = UNORDERED_RE.match(line)
        ordered = ORDERED_RE.match(line)
        if unordered or ordered:
            flush_paragraph(doc, paragraph_buffer)
            if quote_buffer:
                add_quote(doc, quote_buffer)
                quote_buffer = []
            if not list_buffer:
                list_ordered = bool(ordered)
            else:
                list_ordered = list_ordered and bool(ordered)
            item = re.sub(r"^\s*(?:[-*+]|\d+\.)\s+", "", line).strip()
            list_buffer.append(item)
            continue

        if quote_buffer:
            add_quote(doc, quote_buffer)
            quote_buffer = []
        if list_buffer:
            for item in list_buffer:
                add_list_item(doc, item, list_ordered)
            list_buffer = []
        paragraph_buffer.append(line)

    flush_all()

    return seen_chapter


def main():
    args = parse_args()
    doc = Document()
    configure_document(doc, args.title, args.author)

    section = doc.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.95)
    section.right_margin = Inches(0.95)

    seen_chapter = False
    for index, md in enumerate(args.inputs):
        seen_chapter = process_markdown(
            doc,
            Path(md),
            args.chapter_level,
            first_file=index == 0,
            seen_chapter=seen_chapter,
        )

    out = Path(args.output)
    doc.save(out)
    print(out)


if __name__ == "__main__":
    main()
